"""Разбор билетов из .docx — ядро, вынесенное из scripts/parse_tests.py.

Чистая функция «файл → строки в схеме таблицы тестов», без Google Sheets и CLI.
Скрипт scripts/parse_tests.py остаётся тонкой обёрткой поверх этого модуля.

Вопросы НЕ генерируются — только структурируется то, что написал преподаватель.

Поверх разбора работает фильтр качества: в тренажёр попадает только целый
вопрос. Всё сомнительное отсекается и складывается в rejected с причиной —
лучше меньше вопросов, чем один кривой у ученика.
"""

from __future__ import annotations

import logging
import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Dict, List, Optional, Tuple

from docx import Document

from services import pdf_tools


logger = logging.getLogger(__name__)

# Версия парсера пишется в манифест. При изменении логики разбора поднять —
# тогда видно, какие файлы разобраны устаревшей версией и требуют перепрогона.
# 2 — добавлено чтение PDF с текстовым слоем. Разбор .docx не изменился,
# поэтому пересобирать накопленное ради версии не требуется.
PARSER_VERSION = 2

# Схема строки — та же, что в таблице тестов проекта. Менять нельзя:
# на неё завязаны хендлеры тестов.
COLS = [
    "Вариант", "Часть", "№", "Вопрос",
    "Вар.1", "Вар.2", "Вар.3", "Вар.4", "Вар.5",
    "Ответ", "Раздел",
]

# А/В — кириллица, A/B — латиница (бывает в docx)
_question_marker_re = re.compile(r"^\s*([АВABаваb])\s*(\d+)\s*[.\)]\s*(.*)$")
_answers_header_re = re.compile(r"^\s*ответы\s*[:\-–—]?\s*$", re.IGNORECASE)
# Колонтитулы страниц: "Вариант 1  1"
_page_header_re = re.compile(r"^Вариант\s+\d+\s+\d+$", re.IGNORECASE)
_part_boundary_re = re.compile(r"^\s*Часть\s+[АБВАВB]\s*$", re.IGNORECASE)

# Задания, которые НЕ разбиваются на варианты
_no_split_re = re.compile(
    r"(определите\s+(верно\s+)?последовательность"
    r"|верно\s+соотнесите"
    r"|установите\s+соответствие"
    r"|соотнесите\s+элементы)",
    re.IGNORECASE,
)

_latin_to_cyr = str.maketrans({"A": "А", "B": "Б", "C": "В", "D": "Г"})

_LETTER_TO_NUM: Dict[str, int] = {
    "А": 1, "Б": 2, "В": 3, "Г": 4, "Д": 5,
    "а": 1, "б": 2, "в": 3, "г": 4, "д": 5,
    "A": 1, "B": 2, "C": 3, "D": 4, "E": 5,
    "a": 1, "b": 2, "c": 3, "d": 4, "e": 5,
}

_option_marker_re = re.compile(
    r"(?:^|(?<=\s)|(?<=;)|(?<=\.)|(?<=\t))"
    r"([1-5АБВГДабвгдABCDEabcde])\s*\)\s*"
)


# ---------- Чтение документа ----------

def _iter_docx_paragraphs(docx_path: Path, keep_tabs: bool = False) -> List[str]:
    doc = Document(str(docx_path))
    out: List[str] = []
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if not keep_tabs:
            t = re.sub(r"\t+", " ", t)
        if t:
            out.append(t)
    return out


def _iter_docx_paragraphs_raw(docx_path: Path) -> List[str]:
    """Сохраняет табы — нужно для двухколоночного формата."""
    return _iter_docx_paragraphs(docx_path, keep_tabs=True)


# ---------- Ключ с ответами ----------

def _normalize_part_b_expected(raw: str) -> str:
    s = (raw or "").strip().upper()
    s = s.replace(" ", " ")
    s = re.sub(r"\s+", "", s)
    s = s.translate(_latin_to_cyr)
    s = re.sub(r"[^АБВГ0-9]", "", s)
    if not s:
        return ""

    only_digits = bool(re.fullmatch(r"\d+", s))
    only_letters = bool(re.fullmatch(r"[АБВГ]+", s))

    if only_digits:
        digits = sorted(set(s), key=lambda x: int(x))
        return "".join(digits)
    if only_letters:
        return s
    return s


def _normalize_expected(raw: str) -> str:
    """Приводит ответ части В к каноническому виду.

    Три разных типа ответа, и путать их нельзя:
      • многовыбор «135» — порядок не важен, сортируем и убираем повторы;
      • последовательность/соответствие «БГВА», «А2Б3В1Г4» — порядок важен;
      • словесный ответ или год «Метрополия», «1795» — оставляем как есть.

    Прежняя версия считала многовыбором любые цифры, из-за чего год 1795
    превращался в 1579, а словесные ответы вычищались в пустую строку.
    """
    s = (raw or "").strip()
    s = s.replace("\xa0", " ")
    s = re.sub(r"\s+", "", s)
    s = s.strip(".,;:")
    if not s:
        return ""

    upper = s.upper().translate(_latin_to_cyr)

    if re.fullmatch(r"\d+", upper):
        # Многовыбор — только цифры вариантов 1–5 без повторов.
        # Годы и прочие числа оставляем нетронутыми.
        digits = list(upper)
        if (
            len(digits) <= 5
            and len(set(digits)) == len(digits)
            and all(d in "12345" for d in digits)
        ):
            return "".join(sorted(digits, key=int))
        return s

    if re.fullmatch(r"[АБВГД]+", upper):
        return upper

    if re.fullmatch(r"(?:[АБВГД]\d){2,}", upper):
        return upper

    # Словесный ответ — сохраняем как написал преподаватель
    return s


# Маркеры частей пишут и кириллицей, и латиницей: «Часть А» / «Часть A»,
# «В1» / «B1». Здесь латинская B — это всегда В (по начертанию), не Б.
_PART_A_CHARS = "АA"
_PART_B_CHARS = "ВB"


def _extract_part_block(text: str, part_chars: str) -> str:
    pattern = (
        rf"Часть\s*[{part_chars}]\s*[:\-–—]\s*(.*?)"
        rf"(?=Часть\s*[{_PART_A_CHARS}{_PART_B_CHARS}]\s*[:\-–—]|$)"
    )
    m = re.search(pattern, text, flags=re.DOTALL | re.IGNORECASE)
    return (m.group(1) if m else "").strip()


def _parse_answers(answer_lines: List[str]) -> Tuple[Dict[str, str], Dict[str, str]]:
    text = "\n".join(answer_lines)

    part_a_block = _extract_part_block(text, _PART_A_CHARS)
    part_b_block = _extract_part_block(text, _PART_B_CHARS)

    answers_a: Dict[str, str] = {}
    answers_b: Dict[str, str] = {}

    source_a = part_a_block or text
    for m in re.finditer(
        rf"[{_PART_A_CHARS}]\s*(\d+)\s*[-–—:]\s*(\d+)", source_a, flags=re.IGNORECASE
    ):
        answers_a[f"А{m.group(1)}"] = m.group(2)

    # Ответ тянется до следующего маркера «Вn —». Разделители между ответами
    # бывают разные: «В1 — Метрополия; В2 — ВАБ» и «В1 — А3Б2В1 В2 — БВГА».
    # Ограничение по маркеру, а не по разделителю, покрывает оба формата
    # и не режет словесные ответы.
    source_b = part_b_block or text
    for m in re.finditer(
        rf"[{_PART_B_CHARS}]\s*(\d+)\s*[-–—:]\s*(.+?)"
        rf"(?=\s*[{_PART_B_CHARS}]\s*\d+\s*[-–—:]|$)",
        source_b,
        flags=re.IGNORECASE | re.DOTALL,
    ):
        answers_b[f"В{m.group(1)}"] = _normalize_expected((m.group(2) or "").strip())

    return answers_a, answers_b


# ---------- Разбор вопросов ----------

@dataclass(slots=True)
class ParsedQuestion:
    part: str  # А/В
    num: str
    question_text: str
    options: List[str]  # 5 слотов
    expected: str


def _split_line_by_numbers(line: str) -> List[Tuple[Optional[int], str]]:
    """Разбивает строку на фрагменты по маркерам N) или А)."""
    matches = list(_option_marker_re.finditer(line))

    if not matches:
        text = line.strip().rstrip(";.,").strip()
        return [(None, text)] if text else []

    result: List[Tuple[Optional[int], str]] = []

    prefix = line[: matches[0].start()].strip().rstrip(";.,\t ").strip()
    if prefix:
        result.append((None, prefix))

    for i, m in enumerate(matches):
        raw_marker = m.group(1)
        num = int(raw_marker) if raw_marker.isdigit() else _LETTER_TO_NUM.get(raw_marker, 0)
        start = m.end()
        end = matches[i + 1].start() if i + 1 < len(matches) else len(line)
        text = line[start:end].strip().rstrip(";.,\t ").strip()
        if 1 <= num <= 5 and text:
            result.append((num, text))

    return result


def _extract_options_from_lines(lines: List[str]) -> Tuple[str, List[str]]:
    """Достаёт варианты ответов из строк после маркера вопроса."""
    if not lines:
        return "", [""] * 5

    opts_map: Dict[int, str] = {}
    unnumbered: List[str] = []

    for ln in lines:
        for num, text in _split_line_by_numbers(ln):
            if num is not None and 1 <= num <= 5:
                opts_map[num] = text
            elif text:
                unnumbered.append(text)

    if opts_map:
        opts = [""] * 5
        for idx, val in opts_map.items():
            opts[idx - 1] = val
        free_slots = [i for i in range(5) if not opts[i]]
        for i, val in enumerate(unnumbered):
            if i < len(free_slots):
                opts[free_slots[i]] = val
        return "", opts

    q_extra = ""
    total = unnumbered
    if len(total) > 5:
        extra_count = len(total) - 5
        q_extra = "\n".join(total[:extra_count])
        total = total[extra_count:]

    opts = [""] * 5
    for i, v in enumerate(total[:5]):
        opts[i] = v

    return q_extra, opts


def parse_paragraphs(docx_path: Path) -> Tuple[List[ParsedQuestion], Dict[str, str], Dict[str, str]]:
    """Основной разбор .docx: вопросы идут абзацами, ключ — после «Ответы»."""
    return parse_lines(_iter_docx_paragraphs_raw(docx_path))


def parse_lines(
    paragraphs_raw: List[str],
) -> Tuple[List[ParsedQuestion], Dict[str, str], Dict[str, str]]:
    """Разбор готовых строк: вопросы идут абзацами, ключ — после «Ответы».

    Отдельно от чтения файла намеренно: строки одинаково приходят из .docx
    и из PDF, и второй парсер для второго формата заводить незачем — иначе
    любая правка разбора чинилась бы дважды.
    """
    paragraphs = [re.sub(r"\t+", " ", t) for t in paragraphs_raw]

    answers_idx: Optional[int] = None
    for i, t in enumerate(paragraphs):
        if _answers_header_re.match(t):
            answers_idx = i
            break

    if answers_idx is None:
        main_lines = paragraphs[::]
        main_lines_raw = paragraphs_raw[::]
        answers_a: Dict[str, str] = {}
        answers_b: Dict[str, str] = {}
    else:
        main_lines = paragraphs[:answers_idx]
        main_lines_raw = paragraphs_raw[:answers_idx]
        answers_a, answers_b = _parse_answers(paragraphs[answers_idx + 1:])

    questions: List[ParsedQuestion] = []
    current_part: Optional[str] = None
    current_num: Optional[str] = None
    current_marker_text: str = ""
    raw_lines_after_marker: List[str] = []
    block_lines: List[str] = []

    def flush() -> None:
        nonlocal current_part, current_num, current_marker_text
        nonlocal raw_lines_after_marker, block_lines

        if not current_part or not current_num:
            return

        if current_part == "А":
            full_marker = current_marker_text.strip()
            if _no_split_re.search(full_marker):
                # Соответствие/последовательность — не дробим на варианты
                all_lines = [full_marker] if full_marker else []
                for ln in raw_lines_after_marker:
                    clean = re.sub(r"\t+", " ", ln).strip()
                    if clean:
                        all_lines.append(clean)
                q_text = "\n".join(all_lines).strip()
                opts = ["", "", "", "", ""]
            else:
                q_extra, opts = _extract_options_from_lines(raw_lines_after_marker)
                q_text = full_marker
                if q_extra:
                    q_text = q_text + "\n" + q_extra if q_text else q_extra
                opts = [o.rstrip(";.,").strip() for o in opts]

            questions.append(ParsedQuestion(
                part="А", num=current_num, question_text=q_text,
                options=opts, expected=answers_a.get(f"А{current_num}", ""),
            ))
        else:
            lines = [current_marker_text.strip()] if current_marker_text.strip() else []
            lines.extend([x for x in block_lines if x.strip()])
            questions.append(ParsedQuestion(
                part="В", num=current_num, question_text="\n".join(lines).strip(),
                options=["", "", "", "", ""], expected=answers_b.get(f"В{current_num}", ""),
            ))

        current_part = None
        current_num = None
        current_marker_text = ""
        raw_lines_after_marker = []
        block_lines = []

    for idx, line in enumerate(main_lines):
        line_raw = main_lines_raw[idx] if idx < len(main_lines_raw) else line

        if _page_header_re.match(line.strip()):
            continue

        if _part_boundary_re.match(line.strip()):
            flush()
            continue

        m = _question_marker_re.match(line)
        if m:
            flush()
            raw_part = m.group(1).upper()
            current_part = {"A": "А", "B": "В", "а": "А", "в": "В", "b": "В", "a": "А"}.get(
                raw_part, raw_part
            )
            current_num = m.group(2)
            current_marker_text = m.group(3) or ""
            raw_lines_after_marker = []
            block_lines = []
            continue

        if not current_part:
            continue

        if current_part == "А":
            raw_lines_after_marker.append(line_raw)
        else:
            block_lines.append(line)

    flush()
    return questions, answers_a, answers_b


def _normalize_part_letter(raw: str) -> str:
    return {"A": "А", "B": "В", "a": "А", "b": "В"}.get(raw, raw)


def parse_tables(docx_path: Path) -> List[ParsedQuestion]:
    """Запасной разбор: вопросы лежат в таблицах Word (формат обобщений).

    Ключа с ответами в таком формате обычно нет — вопросы придут без «Ответа»
    и будут отсеяны фильтром качества. Это ожидаемо и сообщается преподавателю.
    """
    doc = Document(str(docx_path))
    if not doc.tables:
        return []

    questions: List[ParsedQuestion] = []

    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if not cells:
                continue

            m = re.match(r"^\s*([АВABаваb])\s*(\d+)\s*$", cells[0].strip())
            if not m:
                continue

            part = _normalize_part_letter(m.group(1).upper())
            num = m.group(2)

            full_text = cells[-1].strip()
            if not full_text:
                continue

            lines = [ln for ln in full_text.split("\n") if ln.strip()]
            if not lines:
                continue

            if part != "А":
                questions.append(ParsedQuestion(
                    part="В", num=num, question_text="\n".join(lines).strip(),
                    options=["", "", "", "", ""], expected="",
                ))
                continue

            if _no_split_re.search(lines[0]):
                questions.append(ParsedQuestion(
                    part="А", num=num, question_text="\n".join(lines).strip(),
                    options=["", "", "", "", ""], expected="",
                ))
                continue

            q_lines: List[str] = []
            opt_lines: List[str] = []
            found_option = False
            for ln in lines:
                if not found_option and _option_marker_re.search(ln):
                    found_option = True
                (opt_lines if found_option else q_lines).append(ln)

            q_text = "\n".join(q_lines).strip() if q_lines else ""

            if opt_lines:
                opts_map: Dict[int, str] = {}
                unmarked: List[str] = []
                for ln in opt_lines:
                    matches = list(_option_marker_re.finditer(ln))
                    if matches:
                        pre = ln[: matches[0].start()].strip().rstrip(";.,\t ").strip()
                        if pre:
                            unmarked.append(pre)
                        for i_m, match in enumerate(matches):
                            raw = match.group(1)
                            n = int(raw) if raw.isdigit() else _LETTER_TO_NUM.get(raw, 0)
                            start = match.end()
                            end = matches[i_m + 1].start() if i_m + 1 < len(matches) else len(ln)
                            text = ln[start:end].strip().rstrip(";.,\t ").strip()
                            if 1 <= n <= 5 and text:
                                opts_map[n] = text
                    else:
                        clean = ln.strip().rstrip(";.,").strip()
                        if clean and not re.fullmatch(r"[1-5\s\)]+", clean):
                            unmarked.append(clean)

                opts = [opts_map.get(i, "").strip() for i in range(1, 6)]
                free_slots = [i for i in range(5) if not opts[i]]
                for i_u, val in enumerate(unmarked):
                    if i_u < len(free_slots):
                        opts[free_slots[i_u]] = val
            else:
                opts = ["", "", "", "", ""]

            # Маркеров не нашли: первая строка — вопрос, остальные — варианты
            if all(not o for o in opts) and len(lines) > 1:
                q_text = lines[0].strip()
                remaining = [ln.strip().rstrip(";.,") for ln in lines[1:] if ln.strip()]
                opts = [""] * 5
                for i_opt, val in enumerate(remaining[:5]):
                    opts[i_opt] = val

            questions.append(ParsedQuestion(
                part="А", num=num, question_text=q_text, options=opts, expected="",
            ))

    return questions


# ---------- Фильтр качества ----------

MIN_QUESTION_LEN = 10
MIN_OPTIONS = 2

REASON_SHORT_TEXT = "текст вопроса пустой или слишком короткий"
REASON_NO_ANSWER = "не найден правильный ответ"
REASON_BAD_ANSWER = "ответ не совпадает ни с одним вариантом"
REASON_FEW_OPTIONS = "меньше двух вариантов ответа"
REASON_LOST_OPTIONS = "ответ — номер варианта, но сами варианты не разобрались"


def validate(q: ParsedQuestion) -> Optional[str]:
    """Причина отбраковки или None, если вопрос пригоден.

    Тип вопроса определяется наличием вариантов, а не буквой части —
    так же, как это делают хендлеры тестов во время прохождения.
    """
    text = (q.question_text or "").strip()
    if len(text) < MIN_QUESTION_LEN:
        return REASON_SHORT_TEXT

    expected = (q.expected or "").strip()
    if not expected:
        return REASON_NO_ANSWER

    filled = [o for o in q.options if (o or "").strip()]

    if filled:
        # Вопрос с кнопками. Ответом может быть и несколько номеров сразу
        # («135» — многовыбор), движок тестов это поддерживает.
        if not re.fullmatch(r"[1-5]{1,5}", expected):
            return REASON_BAD_ANSWER
        if len(filled) < MIN_OPTIONS:
            return REASON_FEW_OPTIONS
        if any(not (q.options[int(d) - 1] or "").strip() for d in expected):
            return REASON_BAD_ANSWER
    elif q.part == "А" and re.fullmatch(r"[1-5]", expected):
        # У части А варианты обязаны быть. Если их нет, а ответ — номер
        # варианта, значит при разборе они потерялись: ученику не из чего
        # выбирать. Часть В сюда не попадает — там ответ вида «245» пишут
        # цифрами, а перечисление находится в самом тексте вопроса.
        return REASON_LOST_OPTIONS

    return None


def to_row(q: ParsedQuestion, variant: str, section: str = "") -> Dict[str, str]:
    opts = list(q.options) + [""] * (5 - len(q.options))
    return {
        "Вариант": variant,
        "Часть": q.part,
        "№": q.num,
        "Вопрос": q.question_text,
        "Вар.1": opts[0], "Вар.2": opts[1], "Вар.3": opts[2],
        "Вар.4": opts[3], "Вар.5": opts[4],
        "Ответ": q.expected,
        "Раздел": section,
    }


# ---------- Публичный интерфейс ----------

@dataclass
class ParseResult:
    variant: str
    rows: List[Dict[str, str]] = field(default_factory=list)
    rejected: List[Dict[str, str]] = field(default_factory=list)
    source: str = "none"  # paragraphs | tables | none

    @property
    def accepted_count(self) -> int:
        return len(self.rows)

    @property
    def rejected_count(self) -> int:
        return len(self.rejected)

    @property
    def total_found(self) -> int:
        return self.accepted_count + self.rejected_count

    def reasons_summary(self) -> Dict[str, int]:
        out: Dict[str, int] = {}
        for r in self.rejected:
            out[r["reason"]] = out.get(r["reason"], 0) + 1
        return out

    def count_by_type(self) -> Tuple[int, int]:
        """(с вариантами ответа, с текстовым ответом).

        Тип определяется наличием вариантов — так же, как при прохождении теста.
        """
        with_options = sum(
            1 for r in self.rows if any(r[f"Вар.{i}"] for i in range(1, 6))
        )
        return with_options, len(self.rows) - with_options


def parse_docx(path: Path, variant: Optional[str] = None) -> ParseResult:
    """Разбирает один файл — .docx или PDF с текстом.

    Имя оставлено прежним: на него завязаны скрипты и вызовы в боте, а
    формат определяется по расширению, а не по названию функции.

    Разбор таблицами — запасной путь для формата обобщений, и он есть
    только у .docx: в PDF таблица это те же строки текста, их разбирает
    основной путь.
    """
    path = Path(path)
    variant = variant or path.stem
    result = ParseResult(variant=variant)

    if pdf_tools.is_pdf_name(path.name):
        questions, _a, _b = parse_lines(pdf_tools.extract_lines(path))
        source = "pdf" if questions else "none"
        return _finish(result, questions, source, variant, path)

    questions, _a, _b = parse_paragraphs(path)
    source = "paragraphs"

    if not questions:
        questions = parse_tables(path)
        source = "tables" if questions else "none"

    return _finish(result, questions, source, variant, path)


def _finish(
    result: ParseResult,
    questions: List[ParsedQuestion],
    source: str,
    variant: str,
    path: Path,
) -> ParseResult:
    """Фильтр качества и сборка строк — общая для всех форматов."""

    result.source = source

    for q in questions:
        reason = validate(q)
        if reason:
            result.rejected.append({
                "part": q.part,
                "num": q.num,
                "reason": reason,
                "question_text": (q.question_text or "")[:200],
            })
        else:
            result.rows.append(to_row(q, variant))

    logger.info(
        "parse %s: источник=%s принято=%d отбраковано=%d",
        path.name, source, result.accepted_count, result.rejected_count,
    )
    return result


def parse_many(paths: List[Path]) -> List[ParseResult]:
    results: List[ParseResult] = []
    for p in paths:
        try:
            results.append(parse_docx(p))
        except Exception:  # noqa: BLE001
            logger.exception("Не удалось разобрать %s", p)
            failed = ParseResult(variant=Path(p).stem)
            failed.rejected.append({
                "part": "", "num": "",
                "reason": "файл не удалось прочитать",
                "question_text": "",
            })
            results.append(failed)
    return results
